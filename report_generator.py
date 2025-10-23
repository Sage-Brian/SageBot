import re
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from docx import Document

# Register a built-in Unicode font (works on all systems, no TTF file needed)
pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))

# ------------------ TEXT SANITIZATION ------------------
def sanitize_text(text):
    """Replace special characters with safe ASCII equivalents."""
    replacements = {
        '\u2018': "'", '\u2019': "'",  # single quotes
        '\u201c': '"', '\u201d': '"',  # double quotes
        '\u2013': '-', '\u2014': '-',  # dashes
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text

# ------------------ HIGHLIGHTING ------------------
def highlight_phrases(text, highlights):
    """
    Wrap highlighted phrases with <font color='red'> tags for ReportLab.
    """
    for phrase in highlights or []:
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        text = pattern.sub(f"<font color='red'><b>{phrase.upper()}</b></font>", text)
    return text

# ------------------ COLORED PDF GENERATOR ------------------
def generate_colored_pdf(filename, title, sections, highlights=None):
    """
    Create a professional PDF report with color-coded highlights for AI/plagiarism findings.
    """
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=60,
        bottomMargin=60
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Title'],
        fontName='HeiseiMin-W3',
        fontSize=18,
        textColor=colors.darkblue,
        alignment=1,  # Centered
        spaceAfter=20
    )
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Heading2'],
        fontName='HeiseiMin-W3',
        textColor=colors.HexColor("#004c99"),
        spaceBefore=10,
        spaceAfter=8
    )
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='HeiseiMin-W3',
        fontSize=11,
        leading=15
    )

    story = [Paragraph(sanitize_text(title), title_style)]

    for sec_title, sec_text in sections.items():
        story.append(Paragraph(sanitize_text(sec_title), header_style))
        clean_text = sanitize_text(sec_text)
        colored_text = highlight_phrases(clean_text, highlights)
        story.append(Paragraph(colored_text, body_style))
        story.append(Spacer(1, 12))

    doc.build(story)
    return filename

# ------------------ WORD REPORT GENERATOR ------------------
def generate_word_report(filename, title, sections, highlights=None):
    """
    Generate a Word (.docx) report with bold warnings for flagged content.
    """
    doc = Document()
    doc.add_heading(title, 0)

    for sec_title, sec_text in sections.items():
        doc.add_heading(sec_title, level=1)
        paragraph = doc.add_paragraph(sec_text)

        if highlights:
            for phrase in highlights:
                if phrase.lower() in sec_text.lower():
                    paragraph.add_run(f"  [⚠️ {phrase.upper()} DETECTED]").bold = True

    doc.save(filename)
    return filename